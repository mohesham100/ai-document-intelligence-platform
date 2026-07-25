import os
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

def process_uploaded_files(uploaded_files) -> list[Document]:
    """Extracts text from uploaded files and returns LangChain Documents."""
    docs = []
    
    for uploaded_file in uploaded_files:
        file_name = uploaded_file.name
        file_ext = os.path.splitext(file_name)[1].lower()
        
        if file_ext == '.pdf':
            reader = PdfReader(uploaded_file)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    docs.append(Document(
                        page_content=text, 
                        metadata={"source": file_name, "page": i + 1}
                    ))
                    
        elif file_ext == '.docx':
            doc = DocxDocument(uploaded_file)
            text = "\n".join([para.text for para in doc.paragraphs])
            if text.strip():
                docs.append(Document(
                    page_content=text, 
                    metadata={"source": file_name, "page": 1}
                ))
            
        elif file_ext == '.txt':
            text = uploaded_file.getvalue().decode("utf-8")
            if text.strip():
                docs.append(Document(
                    page_content=text, 
                    metadata={"source": file_name, "page": 1}
                ))
            
        elif file_ext == '.csv':
            df = pd.read_csv(uploaded_file)
            text = df.to_string()
            docs.append(Document(
                page_content=text, 
                metadata={"source": file_name, "page": 1}
            ))
            
    # Text splitting configuration based on project requirements
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    
    split_docs = text_splitter.split_documents(docs)
    
    # Inject Chunk IDs
    for i, doc in enumerate(split_docs):
        doc.metadata["chunk_id"] = f"chunk_{i + 1}"   
        
    return split_docs